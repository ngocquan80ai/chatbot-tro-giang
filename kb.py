import os
import io
import pickle

from pypdf import PdfReader       # Thư viện đọc PDF
from docx import Document         # Thư viện đọc DOCX (Word)
import numpy as np               # Thư viện tính toán (dùng để tính độ tương đồng cosine)

from common import create_embedding  # Hàm tạo embedding từ module common

def read_document(file, file_name):
    """
    Đọc nội dung văn bản từ tệp tải lên (PDF, DOCX hoặc TXT).
    - file: đối tượng tệp (mở ở chế độ nhị phân hoặc file-like object, ví dụ UploadedFile của Streamlit).
    - file_name: tên tệp (dùng để nhận diện định dạng dựa trên đuôi file).
    - Trả về: Nội dung văn bản (chuỗi) nếu đọc thành công, hoặc chuỗi rỗng nếu có lỗi.
    """
    # Xác định định dạng tệp dựa vào phần mở rộng của tên file
    ext = os.path.splitext(file_name)[1].lower()  # Lấy đuôi file, chuyển về chữ thường
    text = ""
    try:
        if ext == ".pdf":
            # Đọc file PDF
            reader = PdfReader(file)  # file có thể là một file-like object
            for page in reader.pages:
                # extract_text() trả về nội dung text của trang, thêm newline giữa các trang cho rõ ràng
                text += (page.extract_text() or "") + "\n"
        elif ext == ".docx":
            # Đọc file DOCX
            # Nếu file là bytes, cần dùng BytesIO để Document có thể đọc
            if isinstance(file, (bytes, bytearray)):
                doc = Document(io.BytesIO(file))
            else:
                doc = Document(file)
            # Ghép nội dung tất cả các đoạn (paragraph) lại, ngăn cách bằng newline
            paragraphs = [p.text for p in doc.paragraphs]
            text = "\n".join(paragraphs)
        elif ext == ".txt":
            # Đọc file văn bản thuần
            content = file.read()
            if isinstance(content, bytes):
                # Giải mã bytes sang chuỗi (UTF-8)
                text = content.decode('utf-8', errors='ignore')
            else:
                # Nếu đã là chuỗi (trường hợp file mở sẵn dạng text)
                text = str(content)
        else:
            # Định dạng không hỗ trợ
            return ""
    except Exception as e:
        # In lỗi và trả về chuỗi rỗng nếu đọc không thành công
        print(f"Lỗi khi đọc file {file_name}:", e)
        return ""
    # Làm sạch nội dung: loại bỏ kí tự thừa nếu cần
    text = text.strip()
    return text

def chunk_text(text, max_length=1000):
    """
    Chia văn bản dài thành các đoạn ngắn (chunk) để dễ tạo vector và lưu trữ.
    - text: chuỗi văn bản đầu vào.
    - max_length: độ dài tối đa (tính theo số ký tự) cho mỗi đoạn.
    - Trả về: danh sách các đoạn văn bản (chuỗi), mỗi đoạn có độ dài tối đa max_length.
    """
    # Loại bỏ các khoảng trắng thừa ở đầu/cuối và ký tự xuống dòng kép
    text = text.replace("\r\n", "\n").strip()
    # Tách văn bản thành các từ
    words = text.split()
    chunks = []
    current_chunk = []
    current_length = 0
    for word in words:
        # Thêm +1 cho khoảng trắng khi thêm từ mới (giữa các từ)
        if current_length + len(word) + 1 <= max_length:
            current_chunk.append(word)
            current_length += len(word) + 1
        else:
            # Khi sắp thêm từ vượt quá giới hạn, bắt đầu một chunk mới
            chunks.append(" ".join(current_chunk))
            current_chunk = [word]
            current_length = len(word)
    # Thêm chunk cuối cùng còn sót lại
    if current_chunk:
        chunks.append(" ".join(current_chunk))
    return chunks

def _category_to_filename(category):
    """
    Chuyển tên lớp/chủ đề thành tên tệp tin lưu trữ an toàn.
    """
    # Loại bỏ khoảng trắng ở đầu cuối và thay khoảng trắng bằng dấu gạch dưới
    safe_name = category.strip().replace(" ", "_")
    # Thay thế một số ký tự đặc biệt có thể gây lỗi trong tên file
    for ch in ['/', '\\', ':', '*', '?', '"', '<', '>', '|']:
        safe_name = safe_name.replace(ch, '_')
    filename = f"knowledge_{safe_name}.pkl"
    return filename

def save_vectors(category, texts, vectors):
    """
    Lưu danh sách các vector và đoạn văn bản tương ứng vào tệp theo lớp/chủ đề.
    - category: tên lớp hoặc chủ đề (chuỗi) để phân loại kiến thức.
    - texts: danh sách các đoạn văn bản (chuỗi) đã chia nhỏ từ tài liệu.
    - vectors: danh sách các vector embedding tương ứng với mỗi đoạn trong texts.
    """
    # Xác định tên file để lưu (theo category)
    filename = _category_to_filename(category)
    data = []
    # Nếu file đã tồn tại, đọc dữ liệu cũ để nối thêm (tránh mất dữ liệu cũ)
    if os.path.exists(filename):
        try:
            with open(filename, "rb") as f:
                data = pickle.load(f)
        except Exception as e:
            print("Không thể tải dữ liệu cũ từ", filename, ":", e)
            data = []
    # Ghép cặp vector với đoạn text rồi thêm vào data
    for vec, txt in zip(vectors, texts):
        data.append((vec, txt))
    # Lưu dữ liệu (danh sách các tuple (vector, đoạn văn)) vào file bằng pickle
    with open(filename, "wb") as f:
        pickle.dump(data, f)

def load_vectors(category):
    """
    Tải danh sách các vector và đoạn văn bản đã lưu cho một lớp/chủ đề.
    - category: tên lớp/chủ đề cần tải.
    - Trả về: danh sách các tuple (vector, đoạn văn) nếu có, hoặc [] nếu không có dữ liệu.
    """
    filename = _category_to_filename(category)
    if not os.path.exists(filename):
        return []  # Chưa có file lưu cho chủ đề này
    try:
        with open(filename, "rb") as f:
            data = pickle.load(f)
            return data  # data là list các (vector, text)
    except Exception as e:
        print("Lỗi khi tải dữ liệu từ", filename, ":", e)
        return []

def list_categories():
    """
    Liệt kê danh sách tất cả các lớp/chủ đề đã có dữ liệu lưu trữ.
    Dựa trên các tệp knowledge_*.pkl trong thư mục hiện tại.
    """
    categories = []
    for file in os.listdir("."):
        if file.startswith("knowledge_") and file.endswith(".pkl"):
            name = file[len("knowledge_"):-4]  # cắt bỏ tiền tố và phần mở rộng
            if name:
                # Đổi dấu gạch dưới thành khoảng trắng cho dễ đọc (nếu trước đó có thay thế)
                category_name = name.replace("_", " ")
                categories.append(category_name)
    return categories

def search_question(question, category, top_k=3):
    """
    Tìm kiếm các đoạn kiến thức phù hợp nhất cho câu hỏi trong một lớp/chủ đề.
    - question: câu hỏi (chuỗi văn bản).
    - category: lớp/chủ đề muốn tìm trong kho kiến thức.
    - top_k: số lượng đoạn văn bản liên quan cao nhất muốn lấy ra.
    - Trả về: danh sách tối đa top_k đoạn văn bản liên quan nhất đến câu hỏi.
    """
    # Tạo embedding cho câu hỏi
    q_vector = create_embedding(question)
    if q_vector == [] or q_vector is None:
        return []  # Không tạo được vector cho câu hỏi (có thể lỗi API hoặc question rỗng)
    # Tải dữ liệu kiến thức đã lưu cho chủ đề
    data = load_vectors(category)  # data là list các (vector, text)
    if not data:
        return []  # Nếu chưa có kiến thức nào cho chủ đề này
    # Tính độ tương đồng cosine giữa câu hỏi và từng đoạn kiến thức
    q_vec = np.array(q_vector)
    scores = []
    for vec, text in data:
        vec = np.array(vec)
        # Tính cosine similarity: cos_sim = (q_vec . vec) / (||q_vec|| * ||vec||)
        # Thêm một giá trị rất nhỏ 1e-8 vào mẫu số để tránh chia cho 0
        cosine_score = np.dot(q_vec, vec) / ((np.linalg.norm(q_vec) * np.linalg.norm(vec)) + 1e-8)
        scores.append(cosine_score)
    # Lấy chỉ số của top_k điểm cao nhất
    top_k = min(top_k, len(scores))
    # Sắp xếp các chỉ số dựa theo điểm từ cao đến thấp
    top_indices = sorted(range(len(scores)), key=lambda i: scores[i], reverse=True)[:top_k]
    # Lấy ra các đoạn văn bản tương ứng với các chỉ số trên
    top_texts = [data[i][1] for i in top_indices]
    return top_texts
